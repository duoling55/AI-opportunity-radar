import json
from pathlib import Path

from docx import Document

from opportunity_radar.industry import load_industry_codes
from scripts.import_industry_classification import import_codes


def test_import_codes_writes_deduplicated_json(tmp_path: Path) -> None:
    output = tmp_path / "codes.json"

    import_codes(Path("tests/fixtures/industry_codes.csv"), output)

    assert json.loads(output.read_text(encoding="utf-8")) == [
        {"code": "C", "name": "制造业"},
        {"code": "C34", "name": "通用设备制造业"},
        {"code": "C36", "name": "汽车制造业"},
    ]


def test_load_industry_codes_returns_canonical_code_name_map(tmp_path: Path) -> None:
    output = tmp_path / "codes.json"
    import_codes(Path("tests/fixtures/industry_codes.csv"), output)

    assert load_industry_codes(output) == {
        "C": "制造业",
        "C34": "通用设备制造业",
        "C36": "汽车制造业",
    }


def test_import_codes_qualifies_official_numeric_codes_with_their_section(tmp_path: Path) -> None:
    source = tmp_path / "official.docx"
    document = Document()
    table = document.add_table(rows=4, cols=6)
    table.rows[0].cells[0].text = "门类"
    table.rows[1].cells[0].text = "代码"
    table.rows[2].cells[0].text = "C"
    table.rows[2].cells[4].text = "制造业"
    table.rows[3].cells[1].text = "34"
    table.rows[3].cells[4].text = "通用设备制造业"
    document.save(source)
    output = tmp_path / "codes.json"

    import_codes(source, output)

    assert json.loads(output.read_text(encoding="utf-8")) == [
        {"code": "C", "name": "制造业"},
        {"code": "C34", "name": "通用设备制造业"},
    ]


def test_import_codes_splits_multiple_official_codes_merged_in_one_row(tmp_path: Path) -> None:
    source = tmp_path / "official.docx"
    document = Document()
    table = document.add_table(rows=4, cols=6)
    table.rows[0].cells[0].text = "门类"
    table.rows[1].cells[0].text = "代码"
    table.rows[2].cells[0].text = "C"
    table.rows[2].cells[4].text = "制造业"
    table.rows[3].cells[2].text = "277\n\n278"
    table.rows[3].cells[4].text = "卫生材料制造\n\n药用辅料制造"
    document.save(source)
    output = tmp_path / "codes.json"

    import_codes(source, output)

    assert json.loads(output.read_text(encoding="utf-8")) == [
        {"code": "C", "name": "制造业"},
        {"code": "C277", "name": "卫生材料制造"},
        {"code": "C278", "name": "药用辅料制造"},
    ]


def test_import_codes_keeps_each_hierarchy_level_from_merged_official_rows(
    tmp_path: Path,
) -> None:
    source = tmp_path / "official.docx"
    document = Document()
    table = document.add_table(rows=4, cols=6)
    table.rows[0].cells[0].text = "门类"
    table.rows[1].cells[0].text = "代码"
    table.rows[2].cells[0].text = "C"
    table.rows[2].cells[4].text = "制造业"
    table.rows[3].cells[2].text = "277\n\n278"
    table.rows[3].cells[3].text = "2770\n\n2780"
    table.rows[3].cells[4].text = "卫生材料制造\n\n药用辅料制造"
    document.save(source)
    output = tmp_path / "codes.json"

    import_codes(source, output)

    assert json.loads(output.read_text(encoding="utf-8")) == [
        {"code": "C", "name": "制造业"},
        {"code": "C277", "name": "卫生材料制造"},
        {"code": "C2770", "name": "卫生材料制造"},
        {"code": "C278", "name": "药用辅料制造"},
        {"code": "C2780", "name": "药用辅料制造"},
    ]
